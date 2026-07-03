#!/usr/bin/env python3
"""
Generate System 1 — Model Building: Business Logic & Architecture Overview (DOCX).

Output: docs/notes/content/SYSTEM1_BUSINESS_LOGIC_OVERVIEW.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime, timezone

OUTPUT = __import__("pathlib").Path(__file__).resolve().parent / "SYSTEM1_BUSINESS_LOGIC_OVERVIEW.docx"


def set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def add_subheading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h


def add_code_block(doc, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table with header shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def build():
    doc = Document()

    # ---- TITLE PAGE ----
    title = doc.add_heading("SYSTEM 1 — MODEL BUILDING", level=0)
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('"The Brain" — Offline Intelligence Factory')
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}\n").font.size = Pt(9)
    meta.add_run("Repository: Scalable Brain\n").font.size = Pt(9)
    meta.add_run("Scope: System 1 only (Layers 0–3 + NLP auxiliaries)").font.size = Pt(9)

    doc.add_page_break()

    # ============================================================
    # SECTION 1 — HIGH-LEVEL ARCHITECTURE
    # ============================================================
    add_subheading(doc, "1. High-Level System Architecture", level=1)

    doc.add_paragraph(
        "System 1 is the offline model-building pipeline of the Scalable Brain platform. "
        "It runs on Computer 1 (the training cluster), has no broker access, and places no trades. "
        "Its sole purpose is to transform raw market history and macroeconomic intelligence into "
        "versioned, validated, deployable decision artifacts consumed by other systems."
    )

    add_subheading(doc, "1.1 What System 1 Produces", level=2)

    doc.add_paragraph(
        "System 1 has exactly two outputs — everything else is internal machinery:"
    )

    doc.add_paragraph(
        "1. Model Artifact Bundle — published to object storage (or local filesystem) "
        "for Computer 2 (System 2 — Execution Engine) to pull. Contains: "
        "hmm_model.joblib (regime detector), strategy_weights.json (per-regime allocation), "
        "regime_strategy_map.json (ranked strategies per regime), model_metadata.json, "
        "and latest.json (pointer to the newest valid version). SHA256 checksums on every file.",
        style="List Number"
    )
    doc.add_paragraph(
        "2. Scored Signals — published to the Scored_Signal_Queue for System 3 "
        "(Account Management) to consume. Each message carries the signal, its model score, "
        "the regime context, the approval decision, and the bundle version that produced it.",
        style="List Number"
    )

    add_subheading(doc, "1.2 Pipeline Overview (Top-Down)", level=2)

    doc.add_paragraph(
        "The pipeline runs sequentially, with each step consuming the outputs of previous steps. "
        "Below is the data flow from raw market data to deployable artifacts."
    )

    add_code_block(doc, """
  OANDA v20 API (D1/H4/W1 candles, 2005-present)
        |
        v
  [MODEL-001] Multi-Timeframe Ingestion
        |  Idempotent upserts into Fact_Market_Prices
        |  Data-quality gates, quarantine bad rows
        v
  [MODEL-002] Feature Engineering Pipeline
        |  Trailing indicators (ATR, ADX, returns, vol, price-position)
        |  Versioned Parquet feature store (deterministic, leakage-free)
        v
  +------------------+------------------+
  |                                     |
  v                                     v
  [MODEL-003]                       [MODEL-006]
  Regime Engine HMM                  ML Gatekeeper
  |  4-state Gaussian HMM            |  XGBoost/LightGBM classifier
  |  Probabilistic regime labels     |  Regime-prob features
  |  Persistence smoothing (3 bars)  |  Dynamic per-regime thresholds
  |  K-Means fallback                |  OOS uplift analysis
  v                                     |
  [MODEL-004]                           |
  Per-Regime Attribution                |
  |  Tag every historical trade         |
  |  with entry regime                  |
  |  Win-rate / PF / Sharpe per cell    |
  |  Bayesian shrinkage for thin cells  |
  v                                     |
  [MODEL-005]                           |
  Strategy Vetting & Regime Map         |
  |  6 strict gates per regime          |
  |  Rank qualifying strategies         |
  |  Emit regime_strategy_map.json      |
  |  Emit strategy_weights.json         |
  v                                     |
  [MODEL-007] <-------------------------+
  Model Serializer & Artifact Registry
  |  Bundle + checksum all artifacts
  |  Upload to object storage
  |  Atomic latest.json pointer
  v
  +------------------+------------------+
  |                                     |
  v                                     v
  Object Storage                    Scored_Signal_Queue
  (pulled by System 2)              (consumed by System 3)
        |
        v
  [MODEL-008]  Scored Signal Producer
  [MODEL-009]  Retraining Scheduler (weekly + trigger-based)
  [MODEL-010]  FinBERT Macro Features (optional, P3)
""")

    add_subheading(doc, "1.3 Canonical Granularity Set", level=2)

    doc.add_paragraph(
        "System 1 uses a defined set of timeframes, each with a specific role:"
    )

    styled_table(doc,
        ["Granularity", "Role", "Purpose"],
        [
            ["D1 (Daily)", "Primary", "Modeling and regime detection. Core analytical timeframe."],
            ["H4 (4-Hour)", "Entry", "Trade entry timing. Bridges daily context to execution."],
            ["W1 (Weekly)", "Macro Context", "Long-term trend backdrop. Added by MODEL-001."],
            ["H1 (Hourly)", "Legacy (preserved)", "Historical Layer 2/3 signal and gatekeeper path. Kept for backward compatibility."],
        ]
    )

    doc.add_page_break()

    # ============================================================
    # SECTION 2 — DETAILED MODEL BREAKDOWNS
    # ============================================================
    add_subheading(doc, "2. Detailed Business Logic — Per MODEL", level=1)

    # ----- MODEL-001 -----
    add_subheading(doc, "2.1 MODEL-001 — Multi-Timeframe Data Ingestion", level=2)

    p = doc.add_paragraph()
    p.add_run("Priority: P0-Critical | Effort: 4d | Status: COMPLETE").bold = True

    add_subheading(doc, "Business Purpose", level=3)
    doc.add_paragraph(
        "Build a reliable, resumable pipeline that pulls candlestick price data from the OANDA "
        "v20 REST API across three timeframes (D1 daily, H4 4-hour, W1 weekly) from 2005 to the present. "
        "Every row must be traceable to its ingestion run, every bad row must be quarantined, "
        "and the pipeline must survive interruption without data loss or duplication."
    )

    add_subheading(doc, "Data Sources", level=3)
    styled_table(doc,
        ["Source", "Details"],
        [
            ["OANDA v20 Practice API", "REST endpoint at api-fxpractice.oanda.com. Only practice (paper) account — System 1 never touches a live broker."],
            ["ForexBrainDB (PostgreSQL 16 + TimescaleDB)", "Destination. Connects via src/common/db.py (SQLAlchemy 2.0 + psycopg2). Table: Fact_Market_Prices (TimescaleDB hypertable)."],
        ]
    )

    add_subheading(doc, "What Is Ingested", level=3)
    styled_table(doc,
        ["Field", "Meaning"],
        [
            ["Instrument", "EUR_USD, GBP_USD, USD_JPY, AUD_USD, USD_CAD (5 forex majors)."],
            ["Granularity", "D1 (daily, OANDA code 'D'), H4 (4-hour), W1 (weekly, OANDA code 'W')."],
            ["OHLCV", "Open, High, Low, Close, Volume — the standard candlestick fields. 'Open' and 'Close' are SQL reserved words and are double-quoted in queries."],
            ["Complete", "Boolean: only fully closed candles are ingested. The currently-forming (incomplete) candle is skipped."],
            ["Lineage", "source='OANDA', ingest_run_id (UUID), ingested_at_utc (timestamp) — added to every row for full traceability."],
        ]
    )

    add_subheading(doc, "Business Rules", level=3)
    rules = [
        ("Idempotent Ingestion", "Uses INSERT ... ON CONFLICT (asset_id, granularity, timestamp) DO UPDATE. Running the pipeline twice on the same range produces identical row counts — no duplicates."),
        ("Resumable Cursors", "A cursor (last ingested timestamp) is saved per (instrument, granularity) in results/state/ingest_progress.json. If the pipeline is killed mid-backfill, it resumes from the cursor without re-fetching."),
        ("Rate Limiting", "Chunks requests at 500 candles per call. Exponential backoff with jitter on HTTP 429 (rate limit) and 5xx errors."),
        ("Data Quality Gates", "Pre-commit checks per batch: monotonic timestamps (no out-of-order bars), OHLC sanity (Low <= Open,Close <= High, no negative/zero prices), duplicate detection on natural key."),
        ("Quarantine", "Rows failing any DQ check go to Fact_Market_Prices_Quarantine with a reason code. They are never silently dropped."),
        ("Gap Reporting", "Expected-bar coverage report per instrument/granularity. Weekend and holiday gaps are expected and logged; unexpected gaps (<0.5% missing expected bars) raise a flag."),
        ("Granularity Code Mapping", "OANDA uses 'D' for daily and 'W' for weekly, not 'D1'/'W1'. A to_oanda_granularity() mapper handles this translation. This was a latent bug in earlier code."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    add_subheading(doc, "Current Data Baseline", level=3)
    styled_table(doc,
        ["Granularity", "Rows", "First Bar", "Last Bar", "Note"],
        [
            ["D1", "29,243", "2005-12-31", "2026-04-30", "Primary modeling timeframe"],
            ["H4", "164,563", "2006-01-01", "2026-06-23", "Entry timing"],
            ["W1", "5,340", "2005-12-30", "2026-06-12", "Added by MODEL-001 (1,068 weeks x 5 majors)"],
            ["H1", "648,195", "2006-01-01", "2026-06-23", "Legacy Layer 2/3 path"],
        ]
    )

    # ----- MODEL-002 -----
    add_subheading(doc, "2.2 MODEL-002 — Feature Engineering Pipeline", level=2)

    p = doc.add_paragraph()
    p.add_run("Priority: P1-High | Effort: 3d | Status: COMPLETE").bold = True

    add_subheading(doc, "Business Purpose", level=3)
    doc.add_paragraph(
        "Transform raw OHLCV prices into a canonical set of features stored in a versioned, "
        "deterministic feature store. This eliminates ad-hoc indicator recomputation across layers "
        "and guarantees that every downstream model trains on the same, reproducible features."
    )

    add_subheading(doc, "Feature Definitions", level=3)
    styled_table(doc,
        ["Feature", "Window", "Formula / Meaning"],
        [
            ["returns_1", "1 bar", "Log return of Close over 1 bar: ln(Close_t / Close_t-1). Captures immediate directional movement."],
            ["atr_14", "14 bars", "Average True Range. Measures volatility — the average range between high and low accounting for gaps. Reuses the proven ATR implementation from src/layer0/indicators.py."],
            ["adx_14", "14 bars", "Average Directional Index. Measures trend strength (0–100). Values > 25 indicate a trending market; < 20 indicate ranging. Used as a regime feature input."],
            ["price_position_20", "20 bars", "Where Close sits within the 20-bar high-low channel: (Close - min(Low,20)) / (max(High,20) - min(Low,20)). Output in [0,1]. 0 = at the 20-bar low; 1 = at the 20-bar high. Indicates overbought/oversold relative to recent range."],
            ["volatility_20", "20 bars", "Rolling standard deviation of returns_1 over 20 bars. Measures recent price variability."],
            ["trend_20", "20 bars", "Weighted trend direction (derived, x3 weighting). Helps the HMM distinguish trending-up from trending-down rather than just volatility bands."],
        ]
    )

    add_subheading(doc, "Feature Engineering Rules", level=3)
    rules = [
        ("Point-in-Time (No Look-Ahead)", "All windows are trailing only. No feature at bar t uses data from bar t+1 or later. The first N-1 bars are null (warm-up) and excluded from training. This is the single most important rule — violating it creates false confidence in backtests."),
        ("Determinism", "Running the pipeline twice on identical inputs produces byte-identical Parquet partitions (verified by SHA256). No wall-clock timestamps inside feature columns. Column order is fixed."),
        ("Versioned Storage", "Features are stored as Parquet with Snappy compression, partitioned by granularity and year. Path: feature-store/{version}/granularity={D1|H4|W1}/year=YYYY/part-*.parquet. Each version has a schema.json (column names, dtypes, window params) and lineage.json (source ingest runs, git SHA, build timestamp)."),
        ("Single Source of Truth", "Downstream tasks (MODEL-003 regime, MODEL-006 gatekeeper) read features from here. They do NOT recompute indicators independently."),
        ("Schema Evolution", "New features bump the minor version. Breaking changes (renamed/removed columns) bump major and write to a new path. Old consumers are unaffected."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # ----- MODEL-003 -----
    add_subheading(doc, "2.3 MODEL-003 — Regime Engine HMM Upgrade", level=2)

    p = doc.add_paragraph()
    p.add_run("Priority: P1-High | Effort: 5d | Status: COMPLETE").bold = True

    add_subheading(doc, "Business Purpose", level=3)
    doc.add_paragraph(
        "Detect the current market regime (state) for every bar across all timeframes. "
        "This answers the question: 'Are we trending up, trending down, stuck in a range, "
        "or in a high-volatility breakout?' The answer determines which trading strategies "
        "should be active and how aggressively the gatekeeper should filter signals."
    )

    add_subheading(doc, "The Four Market Regimes", level=3)
    styled_table(doc,
        ["Regime", "Characteristics", "What It Means for Trading"],
        [
            ["Trending-Up", "Persistent positive returns, directional drift upward, above-average ADX.", "Trend-following strategies dominate. Mean-reversion strategies are dangerous (catching a falling knife in reverse)."],
            ["Trending-Down", "Persistent negative returns, directional drift downward.", "Short-biased strategies preferred. Long-only strategies should be suppressed."],
            ["Ranging", "Low directional drift, price oscillating within a channel, low volatility.", "Mean-reversion and range-trading strategies excel. Breakout strategies produce false signals."],
            ["High-Vol (Volatile)", "Elevated volatility, large price swings, high ATR.", "Position sizing must shrink. Wider stops. Breakout strategies may work but with higher risk."],
        ]
    )

    add_subheading(doc, "ML Model: Gaussian Hidden Markov Model (HMM)", level=3)
    doc.add_paragraph(
        "An HMM is a probabilistic sequence model. It assumes the market is always in one of N hidden states "
        "(here, 4 regimes), and each state emits observable data (feature vectors) according to a Gaussian "
        "distribution. The model learns both the state transition probabilities (how likely the market "
        "is to stay in or leave a regime) and the emission distributions (what each regime 'looks like' "
        "in feature space)."
    )

    add_subheading(doc, "HMM Configuration & Training", level=3)
    styled_table(doc,
        ["Parameter", "Value", "Rationale"],
        [
            ["n_components", "4", "One per regime. 4 is validated: fewer loses distinction (e.g., Trending-Up and Trending-Down collapse), more creates noise states."],
            ["Covariance type", "Full", "Full covariance captures relationships between features within each regime. Diagonal is simpler but loses cross-feature interaction."],
            ["Algorithm", "EM (Expectation-Maximization)", "Standard for HMMs. Iteratively estimates state assignments (E-step) and updates parameters (M-step) until log-likelihood converges."],
            ["Random restarts (n_init)", "3", "HMMs are sensitive to initialization. Multiple restarts with a fixed seed (for reproducibility) avoid local optima."],
            ["Convergence tolerance", "1e-4", "Default tolerance. Training stops when log-likelihood improvement drops below this."],
            ["Input features", "atr_14, adx_14, volatility_20, returns_1, trend_20 (weighted x3)", "A mix of volatility and directional features. trend_20 is weighted x3 to help the HMM learn direction (not just volatility bands)."],
        ]
    )

    add_subheading(doc, "Key Business Rules", level=3)
    rules = [
        ("Deterministic State-to-Label Mapping", "HMM states (0,1,2,3) are unlabeled — they have no inherent meaning. After fitting, states are mapped to semantic labels by their component means: highest volatility+ATR → High-Vol, lowest drift+low vol → Ranging, positive mean returns → Trending-Up, negative → Trending-Down. This mapping rule is stored with the model so labels are stable across retrains."),
        ("Probabilistic Outputs", "For every bar, the HMM outputs a full 4-way probability vector (prob_trending_up, prob_trending_down, prob_ranging, prob_high_vol) summing to 1.0. The raw regime label is the argmax (highest probability). Probabilities feed MODEL-006 as gatekeeper features — they tell the model not just 'what regime' but 'how confident'."),
        ("Persistence Smoothing (Min 3 Bars)", "Raw HMM labels can flicker — switching regimes every bar creates noise. A causal persistence filter is applied: a new regime is only accepted once it persists for ≥3 consecutive bars. Otherwise the prior regime is held. This is a debounce, not a future-aware filter. Both regime_raw and regime_smoothed are stored."),
        ("K-Means Fallback", "K-Means clustering is retained as a first-class fallback. If the HMM fails to converge, produces degenerate covariances, or has an effective regime count <4, the system automatically falls back to K-Means and logs the decision. Probabilities for K-Means are one-hot (hard assignment)."),
        ("Stability Gate", "HMM must achieve regime classification accuracy ≥70% on a labeled holdout. Flicker rate (regime switches per bar) must be materially lower than the K-Means baseline. Current stability accuracy: D1=0.886, H4=0.970, H1=0.860."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    add_subheading(doc, "Output Schema (Fact_Market_Regime_V2)", level=3)
    styled_table(doc,
        ["Column", "Type", "Meaning"],
        [
            ["regime_model", "text", "'HMM' or 'KMeans' — which model produced this row."],
            ["regime_raw", "text", "Argmax of posterior probabilities. May flicker."],
            ["regime_smoothed", "text", "Persistence-filtered label (min 3 bars). Use this for downstream decisions."],
            ["prob_trending_up/down/ranging/high_vol", "float", "Four-way probability vector summing to 1.0 per bar."],
            ["model_version", "text", "Version tag (e.g., 'hmm-v1.0.0') for lineage."],
        ]
    )

    # ----- MODEL-004 -----
    add_subheading(doc, "2.4 MODEL-004 — Per-Regime Strategy Attribution", level=2)

    p = doc.add_paragraph()
    p.add_run("Priority: P1-High | Effort: 3d | Status: COMPLETE").bold = True

    add_subheading(doc, "Business Purpose", level=3)
    doc.add_paragraph(
        "A trading strategy that works brilliantly in trending markets may bleed money in ranging markets. "
        "MODEL-004 answers: 'For each strategy, how does it perform in each specific regime?' "
        "This is the foundation for regime-aware strategy selection."
    )

    add_subheading(doc, "Strategy Families Analyzed", level=3)
    styled_table(doc,
        ["Family", "Strategies", "Logic"],
        [
            ["Range Bollinger", "Range_Bollinger, Range_Bollinger_H1, Range_Bollinger_H4, Range_Bollinger_Aggressive", "Mean reversion: buy at lower Bollinger Band with RSI oversold, sell at upper band with RSI overbought. Aggressive variant skips RSI filter."],
            ["VCP Breakout", "VCP_Breakout, VCP_Breakout_H1, VCP_Breakout_H4, VCP_Breakout_Aggressive", "Volatility Contraction Pattern breakout: enters after a squeeze (low volatility) when price breaks the Donchian channel. Trend-aligned entries only."],
            ["Support/Resistance", "Support_Resistance, Support_Resistance_Breakout", "Bounce off S/R levels with RSI confirmation, or breakout through levels. Multi-timeframe context from H4 for H1 entries."],
            ["Trend Donchian", "Trend_Donchian, Trend_Donchian_Squeeze", "Donchian Channel breakout with ADX trend strength filter. Squeeze variant adds Bollinger Band squeeze pre-condition."],
        ]
    )

    doc.add_paragraph(
        "10 strategies total, covering trend-following, mean-reversion, and breakout styles. "
        "66,743 backtested trades were generated via the Layer 0 backtest engine (persist_trade_outcomes.py) "
        "and tagged with their entry regime."
    )

    add_subheading(doc, "Attribution Methodology", level=3)
    rules = [
        ("Point-in-Time Regime Tagging", "Each trade is tagged with the regime_smoothed label active at or just before the trade's entry timestamp, using merge_asof (backward nearest join). The regime label at entry is all that was knowable at the time. No future regime is used."),
        ("Per-Cell Metrics", "For each combination of strategy x regime x granularity (80 cells total), compute: win_rate (wins/total), profit_factor (gross profit/gross loss), sharpe_ratio (annualized), expectancy (average R per trade), max_drawdown, avg_R (average return per trade), and trade_count."),
        ("Bayesian Shrinkage (N_min = 20)", "Cells with fewer than 20 trades are flagged low_confidence=true. Their metrics are shrunk toward the strategy's global (aggregate) metrics using a Bayesian prior. This prevents a strategy from looking amazing in a regime where it only took 3 lucky trades."),
        ("Reconciliation", "The sum of per-regime trade counts must equal the aggregate trade count. Zero cells indicate a strategy was never active during that regime — they are absent, not zero-padded."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    add_subheading(doc, "Attribution Results (Current)", level=3)
    styled_table(doc,
        ["Regime", "Trades", "Share", "Interpretation"],
        [
            ["Ranging", "37,944", "56.9%", "Most trades occurred in ranging markets — consistent with mean-reversion strategies dominating."],
            ["Trending-Down", "15,592", "23.4%", "Significant exposure to downtrends. Short-biased strategies captured these."],
            ["Trending-Up", "9,140", "13.7%", "Fewer trend-up trades. Trend-following long strategies had limited opportunities."],
            ["High-Vol", "4,067", "6.1%", "Sparse regime — few trades in extreme volatility. High shrinkage applied."],
        ]
    )

    # ----- MODEL-005 -----
    add_subheading(doc, "2.5 MODEL-005 — Strategy Vetting & Regime Map", level=2)

    p = doc.add_paragraph()
    p.add_run("Priority: P1-High | Effort: 3d | Status: READY TO START").bold = True
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0x55, 0x00)

    add_subheading(doc, "Business Purpose", level=3)
    doc.add_paragraph(
        "Apply strict quality gates to each strategy in each regime, then rank the survivors. "
        "The output answers: 'If the market is in Regime X, which strategies should I trust, "
        "and in what order?' The map and weights file are a hard interface contract to Computer 2."
    )

    add_subheading(doc, "The Six Vetting Gates", level=3)
    doc.add_paragraph(
        "Each strategy must pass ALL six gates per regime (with adequate sample size) to qualify. "
        "These metrics were chosen because together they paint a complete picture of risk-adjusted "
        "performance, not just raw profitability."
    )

    styled_table(doc,
        ["Gate", "Threshold", "Why This Metric"],
        [
            ["Profit Factor", ">= 1.5", "Gross Profit / Gross Loss. A PF of 1.5 means for every $1 lost, $1.50 is won. Above 1.0 is profitable; 1.5 ensures a comfortable margin. Chosen because it is the most intuitive profitability ratio — immune to position sizing."],
            ["Sharpe Ratio", ">= 0.8", "Annualized (mean return / std of returns). Measures risk-adjusted return. 0.8 means you earn 0.8 units of return per unit of risk. Chosen because it penalizes volatile return streams — a strategy with high returns but wild swings fails."],
            ["Max Drawdown", "<= 25%", "Largest peak-to-trough decline. A 25% drawdown means at some point the account was down 25% from its peak. Chosen because drawdowns kill accounts — even a profitable strategy is dangerous if it first loses a quarter of capital."],
            ["Win Rate", ">= 40%", "Winning trades / total trades. A 40% win rate means 4 out of 10 trades win. Chosen because low win rates (e.g., 25%) cause psychological abandonment — traders (and automated systems monitoring drawdowns) exit during losing streaks."],
            ["Recovery Factor", ">= 3.0", "Net Profit / Max Drawdown. Recovery of 3.0 means the strategy recoups its worst drawdown 3x over. Chosen because it measures resilience — can the strategy climb out of its hole?"],
            ["OOS Coverage", ">= 60 months", "The strategy must have demonstrated performance over at least 60 months of out-of-sample data (walk-forward folds). Chosen because 5 years covers multiple market cycles. Prevents strategies that look good on a 6-month lucky window."],
        ]
    )

    add_subheading(doc, "Vetting Rules", level=3)
    rules = [
        ("All-or-Nothing per Gate", "A strategy that passes 5 of 6 gates but fails MaxDD still fails entirely. No compensatory logic — a single broken gate disqualifies."),
        ("Low-Confidence Cells Cannot Qualify", "Cells flagged low_confidence (fewer than N_min=20 trades) in MODEL-004 cannot qualify for that regime regardless of metrics. Bayesian shrinkage prevents lucky small samples from looking good, but the cell is still untrustworthy for promotion."),
        ("Ranking Rule", "Within each regime, qualifying strategies are ranked by a composite score: 0.5*Sharpe + 0.3*PF + 0.2*RecoveryFactor, with a MaxDD penalty. This weights risk-adjusted return highest, profitability next, resilience last. Ties broken by higher trade count (more data = more confidence)."),
        ("Starvation Guard", "If a regime ends up with zero qualifying strategies, the map emits the regime explicitly with an empty list and a warning. No silent omission — downstream systems must know they have no strategies for that regime."),
        ("Log-Only Mode First", "Initial runs compute pass/fail and would-be artifacts without altering production. This allows gate calibration on real data before committing to the stricter regime. Rejection reasons are reported per gate."),
        ("Legacy Compatibility", "The existing aggregate vetting in src/layer0/qualify_strategies.py continues to run. Per-regime vetting is additive. results/sql/layer2_strategies.sql emission is preserved for Layer 2 compatibility."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    add_subheading(doc, "Output Artifacts", level=3)
    styled_table(doc,
        ["Artifact", "Contents", "Consumer"],
        [
            ["regime_strategy_map.json", "Per regime: ranked list of qualifying strategies with metrics, rank, and composite score. Includes schema_version, regime_model_version, ranking_rule, generated_at_utc.", "MODEL-007 (bundle) → System 2 (execution engine)"],
            ["strategy_weights.json", "Per regime: normalized allocation weights over qualifying strategies (sum to 1.0). Derived from composite scores. Includes same versioning fields.", "MODEL-007 (bundle) → System 2"],
        ]
    )

    # ----- MODEL-006 -----
    add_subheading(doc, "2.6 MODEL-006 — ML Gatekeeper: Regime Features & Dynamic Threshold", level=2)

    p = doc.add_paragraph()
    p.add_run("Priority: P2-Medium | Effort: 4d | Status: BLOCKED (needs fact_signals populated via Layer 2 generate_signals)").bold = True
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    add_subheading(doc, "Business Purpose", level=3)
    doc.add_paragraph(
        "The gatekeeper is the final filter. Layer 2 generates trading signals. The gatekeeper scores "
        "each signal (0–1 = 'how good is this opportunity?') and approves/rejects it. MODEL-006 upgrades "
        "the gatekeeper with regime awareness and dynamic thresholds so it can be stricter in unfavorable "
        "regimes and more permissive in favorable ones."
    )

    add_subheading(doc, "ML Models", level=3)
    styled_table(doc,
        ["Algorithm", "Strengths", "Limits"],
        [
            ["XGBoost", "Gradient-boosted trees. Excellent with tabular data. Handles non-linear relationships. Built-in regularization prevents overfitting. Fast training.", "Requires hyperparameter tuning. Can overfit on small datasets. Less interpretable than logistic regression (but SHAP values help)."],
            ["LightGBM", "Similar to XGBoost but faster (leaf-wise tree growth). Better with large datasets. Lower memory usage.", "Can overfit on small datasets (leaf-wise growth). More hyperparameters to tune. Slightly less mature ecosystem."],
            ["RandomForest", "Bagging ensemble. Very robust to noise. Few hyperparameters. Good baseline.", "Slower inference. Less accurate on structured tabular data vs boosting. No built-in early stopping."],
        ]
    )

    doc.add_paragraph(
        "The current gatekeeper runs a tournament: trains XGBoost, LightGBM, and RandomForest, "
        "selects the best on validation metrics, and promotes it as champion. The champion model, "
        "preprocessor (ColumnTransformer with scaling/encoding), and manifest (metadata + SHA256) "
        "are serialized as champion_model.pkl / champion_preprocessor.pkl / champion_manifest.json."
    )

    add_subheading(doc, "New Features (MODEL-006 Addition)", level=3)
    styled_table(doc,
        ["Feature", "Source", "Purpose"],
        [
            ["prob_trending_up", "MODEL-003 HMM", "Probability that current bar is Trending-Up. High value → trend-following signals should get higher scores."],
            ["prob_trending_down", "MODEL-003 HMM", "Probability of Trending-Down. High value → short signals get preference, long signals get penalized."],
            ["prob_ranging", "MODEL-003 HMM", "Probability of Ranging. High value → mean-reversion signals favored."],
            ["prob_high_vol", "MODEL-003 HMM", "Probability of High-Vol. High value → all signals get lower scores (wider stops, higher risk)."],
            ["regime_smoothed", "MODEL-003 HMM", "The persistence-smoothed regime label (one-hot encoded). Allows the model to learn regime-specific scoring patterns."],
            ["regime_confidence", "MODEL-003 HMM", "max(prob_vector). How confident the HMM is in its regime call. Low confidence → conservative scoring."],
        ]
    )

    doc.add_paragraph(
        "All features are joined point-in-time at signal time from Fact_Market_Regime_V2. "
        "They pass through the existing ColumnTransformer/feature-alignment pipeline so train "
        "and inference columns stay aligned."
    )

    add_subheading(doc, "Dynamic Thresholds (Business Logic)", level=3)
    doc.add_paragraph(
        "Today, the gatekeeper uses a single static threshold (LAYER3_APPROVAL_THRESHOLD=0.20) — "
        "any signal scoring ≥0.20 is approved. This is naive: a 0.20 score in a strong uptrend "
        "might be a great opportunity; the same 0.20 score in a high-volatility crash is likely noise."
    )

    doc.add_paragraph("MODEL-006 replaces this with a regime-aware threshold map:")
    add_code_block(doc, """
  {
    "Trending-Up":   0.35,   // Be selective — only the best trend signals
    "Trending-Down": 0.45,   // Even more selective for short signals
    "Ranging":       0.20,   // More permissive — many mean-reversion setups
    "High-Vol":      0.60,   // Very strict — most signals are noise
    "fallback":      0.35    // Global default for unknown/missing regime
  }
""")
    doc.add_paragraph(
        "Thresholds are calibrated per regime on validation folds to maximize OOS P&L "
        "while keeping approval within a turnover band (min 1%, max 35% of signals approved). "
        "A global fallback covers regimes with too few samples to calibrate."
    )

    add_subheading(doc, "OOS Uplift Analysis", level=3)
    doc.add_paragraph(
        "The core question: 'Do approved signals actually make money compared to rejected ones?'"
    )
    rules = [
        ("Walk-Forward Folds", "Data is split into sequential train/validation/test windows. The model is trained on past data and tested on future data — exactly as it would be deployed. Multiple folds ensure robustness."),
        ("Uplift Metric", "P&L of approved signals vs P&L of rejected signals on OOS folds. If approved signals lose money or barely outperform random, the gatekeeper is not adding value."),
        ("Significance Gate", "Uplift must be statistically significant (p < 0.05 via bootstrap or t-test on per-trade returns). This prevents promoting a model that got lucky on a small OOS window."),
        ("Deployment Gate", "A new gatekeeper must show non-negative, significant OOS uplift vs the incumbent before promotion. This is enforced by MODEL-009."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # ----- MODEL-007 -----
    add_subheading(doc, "2.7 MODEL-007 — Model Serializer & Artifact Registry", level=2)

    p = doc.add_paragraph()
    p.add_run("Priority: P0-Critical | Effort: 3d | Status: NOT STARTED (blocked by MODEL-005)").bold = True
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    add_subheading(doc, "Business Purpose", level=3)
    doc.add_paragraph(
        "Package all System 1 outputs into a single, checksummed, versioned bundle and publish it "
        "to object storage. This is the decoupling boundary: Computer 2 (System 2) pulls the bundle "
        "and Computer 1 never talks to the broker. The bundle is the 'brain' that gets deployed."
    )

    add_subheading(doc, "Bundle Contents", level=3)
    styled_table(doc,
        ["File", "Content", "Source Task"],
        [
            ["hmm_model.joblib", "Fitted Gaussian HMM + scaler + state-to-label mapping + per-granularity model weights.", "MODEL-003"],
            ["strategy_weights.json", "Per-regime allocation weights (sum to 1.0 per regime).", "MODEL-005"],
            ["regime_strategy_map.json", "Per-regime ranked list of qualifying strategies with metrics.", "MODEL-005"],
            ["model_metadata.json", "Manifest: bundle_version, schema_version, regime_model_version, feature_set_version, vetting_run_id, mlflow_run_id, per-file SHA256 + byte counts, summary metrics.", "MODEL-007 (generated)"],
            ["checksums.sha256", "SHA256 hash of every artifact (also embedded in metadata).", "MODEL-007 (generated)"],
            ["latest.json", "Pointer: {bundle_version, path, metadata_sha256, promoted_at_utc}. Updated atomically only after full upload + verification.", "MODEL-007 (generated)"],
        ]
    )

    add_subheading(doc, "Publishing Protocol (Security Rules)", level=3)
    rules = [
        ("Compute First, Upload After", "SHA256 is computed locally for every artifact BEFORE upload. model_metadata.json and checksums.sha256 are written locally first."),
        ("Verify After Upload", "After uploading to object storage, every stored object is re-downloaded (or head-requested) and its SHA256 is verified against the local hash. On ANY mismatch: delete the partial upload, abort, do NOT advance latest.json."),
        ("Atomic Pointer Update", "latest.json is written to a temp file and atomically renamed (os.replace on POSIX; precondition check on GCS). A reader never sees a half-written pointer or a pointer to an incomplete bundle."),
        ("Guards Against Bad Bundles", "The serializer REFUSES to promote if: any artifact is missing, any SHA256 mismatches, the regime strategy map is empty (no qualifying strategies for any regime), or upstream quality gates failed. Returns non-zero exit code."),
        ("Encryption", "Storage backend enforces encryption at rest (default on GCS; flagged as N/A for local filesystem during development). TLS for all transfers. No secret (API key, password, credential) is ever serialized into any metadata or artifact."),
        ("Lifecycle / Rollback", "Keep last N (default 5) bundle versions. Older versions are pruned. Rollback = repoint latest.json to the previous good bundle_version. Because Computer 2 verifies checksums before loading, a single pointer revert fully rolls back."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    add_subheading(doc, "Storage Abstraction", level=3)
    doc.add_paragraph(
        "MODEL-007 codes against a pluggable StorageBackend interface (src/common/storage/). "
        "The default is a local-filesystem backend that faithfully reproduces production semantics "
        "(immutable versions, atomic pointer, SHA256 round-trip). Switching to Google Cloud Storage "
        "later is a single .env change (STORAGE_PROVIDER=gcs) with no code change."
    )

    # ----- MODEL-008 -----
    add_subheading(doc, "2.8 MODEL-008 — Scored Signal Queue Producer", level=2)

    p = doc.add_paragraph()
    p.add_run("Priority: P0-Critical | Effort: 2d | Status: COMPLETE").bold = True

    add_subheading(doc, "Business Purpose", level=3)
    doc.add_paragraph(
        "Decouple the training system from execution. Instead of Layer 3 directly calling Layer 4 "
        "(which binds scoring to the live pipeline), System 1 publishes scored signals to a durable "
        "message queue. System 3 (Account Management) consumes them independently. This structural "
        "cut enables Computer 1 (training) and Computer 2 (execution) to operate on different machines "
        "at different cadences."
    )

    add_subheading(doc, "Message Contract", level=3)
    doc.add_paragraph("Every message in the Scored_Signal_Queue carries:")
    styled_table(doc,
        ["Field", "Meaning"],
        [
            ["message_id", "Idempotency key: <signal_id> + <score_run_id>. Publishing the same scored signal twice is a consumer-side no-op."],
            ["signal_id", "Identifier of the signal from Layer 2."],
            ["instrument", "EUR_USD, GBP_USD, etc."],
            ["granularity", "H1 or H4 (legacy contract preserved)."],
            ["signal_time_utc", "When the signal fired (bar timestamp)."],
            ["direction", "'long' or 'short'."],
            ["model_score", "Gatekeeper score in [0.0, 1.0]."],
            ["approved", "True if model_score >= threshold_applied."],
            ["threshold_applied", "The regime-aware threshold used for this decision (from MODEL-006 dynamic threshold map)."],
            ["regime", "Trending-Up / Trending-Down / Ranging / High-Vol."],
            ["regime_probs", "Full 4-way probability vector from MODEL-003 HMM."],
            ["bundle_version", "Link to MODEL-007 bundle that produced this score."],
            ["produced_at_utc", "Timestamp when scored."],
        ]
    )

    add_subheading(doc, "Producer Semantics (Safety Rules)", level=3)
    rules = [
        ("Never Drop Valid Signals", "When the queue is at capacity (MAX_QUEUE_SIZE), the producer applies backpressure — it blocks and retries with exponential backoff rather than dropping messages. A scored signal is a derived trading decision; losing it silently is unacceptable."),
        ("Dead-Letter Queue (DLQ)", "Messages that persistently fail to publish (broker down, serialization error, oversized message) go to a DLQ with a reason code and timestamp. DLQ growth triggers an alert. No message is lost — it either publishes or lands in the DLQ."),
        ("Idempotency", "The message_id is derived deterministically from (signal_id, score_run_id). Re-publishing the same scored signal has no effect — the consumer deduplicates on message_id."),
        ("At-Least-Once Delivery", "Publisher waits for broker confirmation before considering a message published. Combined with consumer-side idempotency, this gives effectively exactly-once semantics."),
        ("Zero Layer-4 Imports", "System 1 has no import, call, or reference to src/layer4_executor/. This is verified by static analysis. The decoupling is structural."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # ----- MODEL-009 -----
    add_subheading(doc, "2.9 MODEL-009 — Retraining Scheduler", level=2)

    p = doc.add_paragraph()
    p.add_run("Priority: P2-Medium | Effort: 3d | Status: NOT STARTED (blocked by MODEL-007)").bold = True
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    add_subheading(doc, "Business Purpose", level=3)
    doc.add_paragraph(
        "Markets change. Strategies that worked last quarter may fail this quarter. Regime behavior "
        "evolves. The brain must retrain on fresh data. MODEL-009 automates this: it runs the full "
        "MODEL-001→007 pipeline on schedule and on performance triggers, but ONLY promotes a new "
        "model if it demonstrably beats the incumbent."
    )

    add_subheading(doc, "Retraining Triggers", level=3)
    styled_table(doc,
        ["Trigger", "Condition", "Rationale"],
        [
            ["Weekly Schedule", "Sunday 00:00 UTC", "Weekend window before markets reopen. Sized to finish within the compute budget (target: a few hours). The weekly cadence ensures the brain never goes stale."],
            ["Rolling 14-day Sharpe < 0.3", "Live performance: annualized Sharpe over last 14 trading days drops below 0.3.", "A period of sustained poor performance. 14 days is long enough to not fire on a single bad day, short enough to catch degradation early. Sharpe < 0.3 means returns are barely above risk-free rate."],
            ["Regime Accuracy < 70%", "HMM predicted regime vs realized regime over recent window falls below 70%.", "The regime model is no longer matching reality. Regimes may have shifted (e.g., a new volatility regime post-crisis). The model needs fresh data to re-learn."],
            ["Circuit Breaker", "Signal from System 2/3 (drawdown breach, anomaly, risk limit hit).", "Emergency trigger. Something is wrong with the current model. Retrain immediately."],
        ]
    )

    add_subheading(doc, "Deployment Gates (Must Pass Before Promotion)", level=3)
    doc.add_paragraph(
        "The retrained model is a CANDIDATE. It does not automatically become the new brain. "
        "It must pass these gates:"
    )
    rules = [
        ("HMM Quality", "HMM converges + regime accuracy ≥ 70%. Ensures the regime detector is still working."),
        ("Vetting Output", "regime_strategy_map.json is non-empty. At least one strategy qualifies in at least one regime. An empty map starves the execution engine."),
        ("Gatekeeper OOS Uplift", "Uplift is non-negative, statistically significant (p<0.05), AND ≥ incumbent's uplift. The new gatekeeper must not be worse than what's already deployed."),
        ("Bundle Integrity", "All SHA256 checksums verify (MODEL-007 round-trip). No corrupted artifact gets promoted."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    add_subheading(doc, "Safety Mechanisms", level=3)
    rules = [
        ("Cooldown Debounce", "A retrain fired within a cooldown window (e.g., 6 hours) of the previous retrain is suppressed. Prevents retrain storms from oscillating triggers."),
        ("Single-Flight Lock", "If a scheduled retrain is running and a performance trigger fires, the trigger is queued/dropped — never run two retrains concurrently."),
        ("Shadow/Canary Mode (Optional)", "A passing candidate can be scored in shadow (publish scores but don't update latest.json) for one cycle before promotion. Human or automated review of the shadow period before the pointer flip."),
        ("One-Click Rollback", "Rollback = repoint latest.json to the previous bundle version. No complex procedure. The scheduler records every promote/skip decision with lineage in MLflow and results/state/."),
        ("Incumbent Protection", "The auto-promotion of a worse champion is the #1 risk. The must-beat-incumbent OOS gate + optional shadow/canary + pointer rollback are the three-layer defense."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    # ----- MODEL-010 -----
    add_subheading(doc, "2.10 MODEL-010 — FinBERT Macro Feature Integration", level=2)

    p = doc.add_paragraph()
    p.add_run("Priority: P3-Low (Optional) | Effort: 3d | Status: NOT STARTED (blocked by MODEL-006)").bold = True
    p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    add_subheading(doc, "Business Purpose", level=3)
    doc.add_paragraph(
        "Central bank announcements, Fed minutes, ECB statements, and economic calendar events move "
        "forex markets. MODEL-010 integrates FinBERT (a financial-domain BERT NLP model) to extract "
        "sentiment from these events and feed it as features to the gatekeeper, plus emit time-based "
        "veto windows (e.g., 'don't trade for 2 hours around FOMC'). This is the lowest priority task "
        "and will only be kept if it demonstrably improves OOS uplift."
    )

    add_subheading(doc, "Current State", level=3)
    doc.add_paragraph(
        "src/nlp/finbert.py and src/nlp/macro_scraper.py already scrape ECB/Fed RSS feeds and "
        "economic calendar events into Fact_Macro_Events. FinBERT scores the text. But this data "
        "is idle — it is NOT used as gatekeeper features and NOT enforced as a veto. It is a "
        "built-but-unconnected auxiliary."
    )

    add_subheading(doc, "Macro Features", level=3)
    styled_table(doc,
        ["Feature", "Source", "Description"],
        [
            ["macro_sentiment_score", "FinBERT on event text", "Polarity score in [-1, 1]. -1 = strongly negative (hawkish, tightening), +1 = strongly positive (dovish, accommodative)."],
            ["macro_event_impact", "Economic calendar", "high / medium / low. FOMC, NFP, ECB rate decisions are high-impact. Minor speeches are low-impact."],
            ["time_to_next_event", "Economic calendar", "Hours until the next scheduled high-impact event. Small value → be cautious (event approaching)."],
            ["time_since_last_event", "Economic calendar", "Hours since last high-impact event. Small value → market may still be reacting to news."],
            ["in_event_window", "Computed", "Boolean: are we currently inside a high-impact event window? True → suppress trading."],
        ]
    )

    add_subheading(doc, "Veto Windows (For System 3)", level=3)
    doc.add_paragraph(
        "System 1 emits advisory veto windows — time intervals around high-impact events where "
        "trading should be suppressed or down-weighted. System 3 enforces the veto. System 1 only "
        "produces the signal; it never blocks execution directly (preserving the decoupling contract)."
    )

    add_subheading(doc, "Key Safety Rule: No Result Look-Ahead", level=3)
    doc.add_paragraph(
        "The schedule of a known upcoming event (e.g., 'FOMC statement at 14:00 ET') is fair game "
        "for features — the market knows the event is coming and prices the uncertainty. But the "
        "RESULT/SENTIMENT of that event is NOT available until the release time. Features joined at "
        "signal time must only use events whose sentiment was published at or before the signal "
        "timestamp. This is the same point-in-time principle as MODEL-002/004."
    )

    add_subheading(doc, "Optional & Removable", level=3)
    doc.add_paragraph(
        "MODEL-010 is behind a feature flag. If adding macro features degrades OOS uplift (the model "
        "overfits to news sentiment that doesn't translate to trade outcomes), the features are "
        "removed. The veto export can remain independently even without model features — a trading "
        "halt around FOMC is a reasonable risk-control measure regardless of sentiment scoring."
    )

    doc.add_page_break()

    # ============================================================
    # SECTION 3 — STATUS SUMMARY
    # ============================================================
    add_subheading(doc, "3. Current Implementation Status", level=1)

    styled_table(doc,
        ["Task", "Priority", "Status", "Blocked By"],
        [
            ["MODEL-001 Multi-TF Ingestion", "P0", "COMPLETE", "—"],
            ["MODEL-002 Feature Store", "P1", "COMPLETE", "—"],
            ["MODEL-003 Regime HMM", "P1", "COMPLETE", "—"],
            ["MODEL-004 Per-Regime Attribution", "P1", "COMPLETE", "—"],
            ["MODEL-005 Strategy Vetting", "P1", "READY", "—"],
            ["MODEL-006 Gatekeeper Upgrades", "P2", "BLOCKED", "fact_signals empty"],
            ["MODEL-007 Serializer/Registry", "P0", "BLOCKED", "MODEL-005"],
            ["MODEL-008 Queue Producer", "P0", "COMPLETE", "—"],
            ["MODEL-009 Retraining Scheduler", "P2", "BLOCKED", "MODEL-007"],
            ["MODEL-010 FinBERT Macro", "P3", "BLOCKED", "MODEL-006"],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Next action: ").bold = True
    p.add_run("Execute MODEL-005 (Strategy Vetting & Regime Map). It is the only unblocked task and "
              "is the critical path dependency for MODEL-007 and MODEL-009.")

    add_subheading(doc, "Critical Path", level=2)
    doc.add_paragraph(
        "MODEL-001 → MODEL-002 → MODEL-003 → MODEL-004 → MODEL-005 → MODEL-007 → MODEL-009"
    )
    doc.add_paragraph(
        "Parallel tracks: MODEL-006 (after MODEL-003, once fact_signals is populated) and "
        "MODEL-008 (complete, independent). MODEL-010 is optional and last."
    )

    doc.add_page_break()

    # ============================================================
    # SECTION 4 — GLOSSARY
    # ============================================================
    add_subheading(doc, "4. Glossary of Key Terms", level=1)

    styled_table(doc,
        ["Term", "Definition"],
        [
            ["Granularity", "The timeframe of a candlestick bar. D1 = 1 day, H4 = 4 hours, W1 = 1 week, H1 = 1 hour."],
            ["Regime", "The current market state or 'weather.' One of: Trending-Up, Trending-Down, Ranging, High-Vol."],
            ["HMM (Hidden Markov Model)", "A probabilistic model that learns hidden states from sequential data. Used for regime detection because markets are sequential — yesterday's regime influences today's."],
            ["EM (Expectation-Maximization)", "The training algorithm for HMMs. Iteratively guesses state assignments and refines parameters."],
            ["K-Means", "A clustering algorithm that groups data into K clusters by minimizing distance to centroids. Simple, fast, deterministic. Retained as HMM fallback."],
            ["Persistence Smoothing", "A debounce filter: a regime change is only accepted after it persists for N consecutive bars. Suppresses flickering labels."],
            ["Point-in-Time (PIT)", "The rule that no model input at time t may use data from time t+1 or later. The single most important data-science discipline in trading — violation = look-ahead bias = fake backtest results."],
            ["Look-Ahead Bias", "Using future data to make a past decision. Example: using the entire day's close to decide an entry at 10am. Produces impossibly good backtests that fail in live trading."],
            ["OOS (Out-of-Sample)", "Data not seen during training. The true test of a model. A model that works in-sample but fails OOS is overfit."],
            ["PF (Profit Factor)", "Gross Profit / Gross Loss. Ratio of money made to money lost. >1.0 is profitable. 1.50 means $1.50 won per $1 lost."],
            ["Sharpe Ratio", "Risk-adjusted return. (Mean Return - RiskFreeRate) / StdDev(Returns). Higher = better return per unit of risk. <0.5 is poor, >1.0 is good, >2.0 is excellent."],
            ["MaxDD (Maximum Drawdown)", "Largest peak-to-trough decline in account equity. Expressed as a percentage. A strategy with 70% returns and 60% MaxDD is dangerous."],
            ["Win Rate", "Winning trades / Total trades. Psychological sustainability metric. Low win rates cause abandonment."],
            ["Recovery Factor", "Net Profit / Max Drawdown. How many times the strategy recoups its worst loss. ≥3.0 means it recovers 3x."],
            ["Bayesian Shrinkage", "Pulling extreme estimates from small samples toward a prior (global) estimate. Prevents a strategy from looking amazing in a regime based on 3 lucky trades."],
            ["ColumnTransformer", "Scikit-learn pipeline component that applies different preprocessing (scaling, one-hot encoding) to different column subsets. Ensures train/inference column alignment."],
            ["Walk-Forward", "Validation method: train on period 1, test on period 2; then train on period 2, test on period 3; etc. Simulates how the model would perform if deployed sequentially."],
            ["Backpressure", "Flow control: when a queue is full, the producer slows down or blocks rather than overwhelming the consumer or dropping messages."],
            ["DLQ (Dead-Letter Queue)", "Destination for messages that could not be processed after retries. Prevents data loss."],
            ["Idempotency", "An operation that produces the same result whether executed once or many times. Critical for resumable pipelines and message queues."],
            ["Atomic Pointer", "A pointer (latest.json) that is updated in a single indivisible operation. A reader always sees either the old or new value, never a half-written one."],
            ["MLflow", "Experiment tracking platform. Records parameters, metrics, artifacts, and model versions across training runs for reproducibility and comparison."],
            ["FinBERT", "A BERT (Bidirectional Encoder Representations from Transformers) model fine-tuned on financial text. Classifies sentiment of financial news."],
            ["Veto Window", "A time interval around a high-impact event during which trading is advisory-suppressed. System 1 emits; System 3 enforces."],
        ]
    )

    # ---- FOOTER ----
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "Generated from docs/implementation-roadmap/system-1-model-building/ source-of-truth documents. "
        "Do not hand-edit numbers — regenerate from source specs."
    )
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
