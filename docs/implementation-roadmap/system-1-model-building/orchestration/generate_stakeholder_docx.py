#!/usr/bin/env python3
"""
================================================================================
Stakeholder DOCX generator — System 1 (Model Building)
================================================================================
Renders ``STAKEHOLDER_UPDATE.docx`` from the machine-readable
``progress_ledger.json`` so stakeholders always see the current high-level
status without anyone hand-editing numbers.

The ``ledger-keeper-agent`` (or the single executing LLM, in a no-sub-agent
runtime) runs this after every state change:

    python docs/implementation-roadmap/system-1-model-building/orchestration/generate_stakeholder_docx.py

Design notes:
- Single source of truth = ``progress_ledger.json`` (sibling file). Percentages,
  statuses, and the data baseline are read from it; nothing is invented here.
- Requires ``python-docx`` (declared in requirements.txt). If it is not yet
  installed, the script prints a clear, actionable message and exits non-zero
  instead of crashing — so it is safe to invoke before Phase 0 installs deps.
- Plain-English labels for each MODEL task live in ``TASK_BLURBS`` below so the
  document reads for a non-technical audience.
================================================================================
"""

from __future__ import annotations

import json
import sys
from datetime import datetime, timezone
from pathlib import Path

HERE = Path(__file__).resolve().parent
LEDGER_JSON = HERE / "progress_ledger.json"
OUTPUT_DOCX = HERE / "STAKEHOLDER_UPDATE.docx"

# Plain-English, stakeholder-facing one-liners per task id.
TASK_BLURBS = {
    "MODEL-001": "Add weekly bars + data-quality checks to existing price loading",
    "MODEL-002": "Turn prices into reproducible, versioned model inputs",
    "MODEL-003": "Detect market state (trending/ranging/volatile) with confidence scores",
    "MODEL-004": "Measure each strategy's performance in each market state",
    "MODEL-005": "Keep only strategies that clear strict quality bars; rank per state",
    "MODEL-006": "Use regime confidence to approve/reject signals more intelligently",
    "MODEL-007": "Bundle + checksum + version the brain for safe handoff",
    "MODEL-008": "Publish scored signals for account management (decoupled, safe)",
    "MODEL-009": "Refresh the brain weekly and when performance slips",
    "MODEL-010": "Optional: factor in central-bank/news sentiment + event vetoes",
}

STATUS_LABEL = {
    "not_started": "Not started",
    "in_progress": "In progress",
    "in_review": "In review",
    "rework": "Rework",
    "blocked": "Blocked",
    "done": "Complete",
}

# Status → fraction-complete weighting used to derive an overall percentage when
# the ledger does not supply per-task percent_complete.
STATUS_WEIGHT = {
    "not_started": 0.0,
    "in_progress": 0.4,
    "in_review": 0.8,
    "rework": 0.5,
    "blocked": 0.3,
    "done": 1.0,
}


def load_ledger() -> dict:
    if not LEDGER_JSON.exists():
        sys.exit(f"ERROR: ledger not found at {LEDGER_JSON}")
    with LEDGER_JSON.open(encoding="utf-8") as fh:
        return json.load(fh)


def compute_percent(ledger: dict) -> int:
    tasks = ledger.get("tasks", [])
    if not tasks:
        return int(ledger.get("overall_percent_complete", 0))
    total = 0.0
    for t in tasks:
        pct = t.get("percent_complete")
        if isinstance(pct, (int, float)) and pct > 0:
            total += pct / 100.0
        else:
            total += STATUS_WEIGHT.get(t.get("status", "not_started"), 0.0)
    return round(100.0 * total / len(tasks))


def main() -> int:
    try:
        from docx import Document  # python-docx
        from docx.shared import Pt
    except ImportError:
        print(
            "ERROR: python-docx is not installed. Install it first:\n"
            "    pip install 'python-docx>=1.1.0'\n"
            "(declared in requirements.txt; Phase 0 installs it into the venv).",
            file=sys.stderr,
        )
        return 1

    ledger = load_ledger()
    overall_pct = compute_percent(ledger)
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    doc = Document()
    doc.add_heading("System 1 — Model Building (“The Brain”)", level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"Stakeholder update · generated {now}").italic = True

    p = doc.add_paragraph()
    p.add_run("Overall status: ").bold = True
    p.add_run(f"{ledger.get('overall_status', 'unknown')} — {overall_pct}% complete")

    doc.add_heading("What System 1 delivers", level=1)
    doc.add_paragraph(
        "System 1 is the offline intelligence factory. It turns market history and macro "
        "news into two validated, versioned outputs: a model bundle (the trading brain) for "
        "the execution computer, and a stream of scored trade signals for account management. "
        "It never places trades itself."
    )

    notes = ledger.get("notes")
    if notes:
        doc.add_heading("Current context", level=1)
        doc.add_paragraph(notes)

    # Data baseline (if recorded)
    baseline = ledger.get("data_baseline", {})
    if baseline.get("recorded"):
        doc.add_heading("Price-data baseline", level=1)
        per = baseline.get("per_granularity", {})
        btable = doc.add_table(rows=1, cols=4)
        btable.style = "Light Grid Accent 1"
        hdr = btable.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
            "Granularity", "Rows", "First bar", "Last bar",
        )
        for g, info in per.items():
            row = btable.add_row().cells
            row[0].text = str(g)
            row[1].text = str(info.get("rows", ""))
            row[2].text = str(info.get("first_bar", ""))
            row[3].text = str(info.get("last_bar", ""))

    doc.add_heading("Section-by-section status", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
        "ID", "Purpose", "Status", "% ",
    )
    for t in ledger.get("tasks", []):
        tid = t.get("id", "")
        cells = table.add_row().cells
        cells[0].text = tid
        cells[1].text = TASK_BLURBS.get(tid, t.get("title", ""))
        cells[2].text = STATUS_LABEL.get(t.get("status", ""), t.get("status", ""))
        pct = t.get("percent_complete")
        if not pct:
            pct = round(100 * STATUS_WEIGHT.get(t.get("status", "not_started"), 0.0))
        cells[3].text = f"{pct}%"

    # Most recent events as the "recent milestone" narrative.
    events = ledger.get("events", [])
    if events:
        doc.add_heading("Recent activity", level=1)
        for ev in events[-5:][::-1]:
            line = doc.add_paragraph(style="List Bullet")
            line.add_run(f"{ev.get('ts_utc', '')} — ").bold = True
            line.add_run(ev.get("detail", ev.get("type", "")))

    # Open risks / rework / blocks
    rework = ledger.get("open_rework_directives", [])
    blocked = ledger.get("blocked_agents", [])
    doc.add_heading("Open issues", level=1)
    if not rework and not blocked:
        doc.add_paragraph("None. No open rework directives or blocked agents.")
    else:
        for r in rework:
            doc.add_paragraph(f"Rework: {r}", style="List Bullet")
        for b in blocked:
            doc.add_paragraph(f"Blocked: {b}", style="List Bullet")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(
        "Auto-generated from progress_ledger.json. Do not hand-edit numbers — "
        "update the ledger and regenerate."
    ).italic = True

    doc.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX} ({overall_pct}% complete, "
          f"{len(ledger.get('tasks', []))} tasks).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
